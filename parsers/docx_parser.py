def extract_docx_text(path):
    try:
        from docx import Document

        document = Document(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        table_text = []
        for table in document.tables:
            for row in table.rows:
                table_text.extend(cell.text for cell in row.cells if cell.text.strip())
        return "\n".join(paragraphs + table_text).strip()
    except Exception as exc:
        raise ValueError(f"Could not parse DOCX file: {exc}") from exc

